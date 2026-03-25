"""
Export to dyslexia-friendly DOCX — multilingual.
"""

import os
import time
import re
import pyphen

dic_en = pyphen.Pyphen(lang="en")


def _syllabify_word(word, difficult_words, lang="en"):
    """Return syllabified version if word is difficult."""
    if lang == "en":
        clean = re.sub(r'[^\w]', '', word).lower()
        if clean in difficult_words:
            return dic_en.inserted(word)
    elif lang == "hi":
        core = re.sub(r'[^\u0900-\u097F]', '', word)
        if core in difficult_words:
            try:
                from indicnlp.syllable import syllabifier
                syls = syllabifier.orthographic_syllabify(word, 'hi')
                if syls and len(syls) > 1:
                    return "-".join(syls)
            except Exception:
                pass
    return word


def generate_accessible_docx(
    text,
    difficult_words=None,
    use_syllables=False,
    font_size=20,
    line_spacing=1.8,
    lang="en"
):
    """Generate a dyslexia-friendly DOCX file."""
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_LINE_SPACING

    if difficult_words is None:
        difficult_words = set()

    os.makedirs("downloads", exist_ok=True)
    filename = "downloads/dyslexia_reader_" + str(int(time.time())) + ".docx"

    # Choose font based on language
    if lang == "hi":
        font_name = "Noto Sans Devanagari"
    else:
        font_name = "OpenDyslexic"

    doc = Document()

    for section in doc.sections:
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)

    # Set default font
    style = doc.styles['Normal']
    style.font.name = font_name
    style.font.size = Pt(font_size)

    paragraphs = text.split("\n")

    for para_text in paragraphs:
        if not para_text.strip():
            doc.add_paragraph("")
            continue

        paragraph = doc.add_paragraph()

        para_format = paragraph.paragraph_format
        para_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        para_format.line_spacing = line_spacing
        para_format.space_after = Pt(8)

        words = para_text.split()

        for i, word in enumerate(words):
            # Get root word for difficulty lookup
            if lang == "hi":
                root = re.sub(r'[^\u0900-\u097F]', '', word)
            else:
                root = re.sub(r'[^\w]', '', word).lower()

            # Syllabify if enabled
            if use_syllables:
                display = _syllabify_word(word, difficult_words, lang=lang)
            else:
                display = word

            # Space between words
            if i > 0:
                space_run = paragraph.add_run(" ")
                space_run.font.size = Pt(font_size)
                space_run.font.name = font_name

            # Word run
            run = paragraph.add_run(display)
            run.font.size = Pt(font_size)
            run.font.name = font_name

            # Highlight difficult words
            if root in difficult_words and len(difficult_words) > 0:
                run.bold = True
                run.font.highlight_color = 7  # Yellow

    doc.save(filename)
    return filename